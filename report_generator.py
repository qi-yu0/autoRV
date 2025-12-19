# ==================== report_generator.py ====================
import pandas as pd
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from models import ValidationResult

class ReportGenerator:
    """报告生成器"""
    
    def generate_word_report(self, result: ValidationResult, output_path: str):
        doc = Document()
        self._add_title(doc, "需求完整性验证报告")
        self._add_document_info(doc, result)
        self._add_summary(doc, result)
        self._add_missing_analysis(doc, result)
        self._add_detail_list(doc, result)
        self._add_improvement_suggestions(doc, result)
        doc.save(output_path)
    
    def generate_excel_report(self, result: ValidationResult, output_path: str):
        data = []
        for req in result.requirements_details:
            data.append({
                "需求ID": req.id,
                "需求类型": req.req_type.value,
                "需求描述": req.text[:200],
                "完整性得分": req.completeness_score,
                "缺失要素": ", ".join(req.missing_elements) if req.missing_elements else "无",
                "整改建议": ", ".join(req.improvement_suggestions) if req.improvement_suggestions else "无",
                "所在片段": req.segment_id
            })
        
        df = pd.DataFrame(data)
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='需求明细', index=False)
    
    def _add_title(self, doc: Document, title: str):
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    
    def _add_document_info(self, doc: Document, result: ValidationResult):
        doc.add_heading('一、文档基本信息', level=1)
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        rows = [
            ('文档名称', result.document_name),
            ('文档ID', result.document_id),
            ('验证耗时', f"{result.validation_time:.2f} 秒")
        ]
        
        for i, (key, value) in enumerate(rows):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = str(value)
        doc.add_paragraph()
    
    def _add_summary(self, doc: Document, result: ValidationResult):
        doc.add_heading('二、完整性验证摘要', level=1)
        score_table = doc.add_table(rows=1, cols=3)
        score_table.style = 'Light Grid Accent 1'
        
        cells = score_table.rows[0].cells
        cells[0].text = '总需求数'
        cells[1].text = '完整需求数'
        cells[2].text = '完整性得分'
        
        score_row = score_table.add_row().cells
        score_row[0].text = str(result.total_requirements)
        score_row[1].text = str(result.complete_requirements)
        score_row[2].text = f"{result.completeness_score:.2f}%"
        doc.add_paragraph()
    
    def _add_missing_analysis(self, doc: Document, result: ValidationResult):
        doc.add_heading('三、缺失要素分析', level=1)
        if not result.missing_elements_by_type:
            doc.add_paragraph("未发现缺失要素，需求完整性良好。")
            return
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = '需求类型'
        header_cells[1].text = '缺失要素类型'
        header_cells[2].text = '数量'
        
        for req_type, missing_dict in result.missing_elements_by_type.items():
            for element_type, count in missing_dict.items():
                row_cells = table.add_row().cells
                row_cells[0].text = req_type
                row_cells[1].text = element_type
                row_cells[2].text = str(count)
        doc.add_paragraph()
    
    def _add_detail_list(self, doc: Document, result: ValidationResult):
        doc.add_heading('四、详细缺失清单', level=1)
        incomplete_reqs = [r for r in result.requirements_details if r.missing_elements]
        
        if not incomplete_reqs:
            doc.add_paragraph("无缺失要素需求。")
            return
        
        for i, req in enumerate(incomplete_reqs[:10], 1):
            doc.add_paragraph(f"{i}. 需求ID: {req.id}", style='List Bullet')
            doc.add_paragraph(f"   描述: {req.text[:100]}...")
            doc.add_paragraph(f"   缺失要素: {', '.join(req.missing_elements)}")
            doc.add_paragraph()
    
    def _add_improvement_suggestions(self, doc: Document, result: ValidationResult):
        doc.add_heading('五、整改建议', level=1)
        all_suggestions = []
        for req in result.requirements_details:
            all_suggestions.extend(req.improvement_suggestions)
        
        unique_suggestions = sorted(list(set(all_suggestions)), key=len, reverse=True)[:10]
        if not unique_suggestions:
            doc.add_paragraph("暂无具体整改建议。")
            return
        
        for i, suggestion in enumerate(unique_suggestions, 1):
            doc.add_paragraph(f"{i}. {suggestion}", style='List Bullet')
        doc.add_paragraph()