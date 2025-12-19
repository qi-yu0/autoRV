"""
需求完整性自动化验证系统演示脚本
文件名：demo_requirement_validator.py
"""

import os
import sys
import json
import time
import shutil
from pathlib import Path
from docx import Document

# 添加当前目录到模块路径
sys.path.append('.')

# 导入模块化组件
from config import Config
from validator import RequirementValidator
from preprocessor import DocumentPreprocessor
from report_generator import ReportGenerator

def print_header(title):
    """打印标题"""
    print("\n" + "="*60)
    print(f"{title:^60}")
    print("="*60)

def create_sample_documents():
    """创建示例需求文档"""
    print_header("创建示例文档")
    
    # 创建示例文档目录
    os.makedirs("sample_docs", exist_ok=True)
    
    # 示例1：智能办公系统需求
    office_content = """
# 智能办公系统需求规格说明书
版本：1.0
日期：2024-01-15

## 1. 项目概述
智能办公系统旨在提升企业内部办公效率，实现流程自动化。

## 2. 功能需求
### 2.1 用户管理模块
2.1.1 用户注册功能
- 新员工可自行注册账号
- 需填写姓名、工号、部门、邮箱等信息
- 密码需包含大小写字母和数字
【缺失：密码复杂度具体要求、注册成功后的邮件通知】

2.1.2 用户登录功能
- 支持用户名/密码登录
- 支持邮箱验证码登录
- 连续失败5次锁定账号30分钟
【完整】

### 2.2 文件管理模块
2.2.1 文件上传功能
- 支持多种格式：doc、docx、pdf、jpg、png
- 单文件最大不超过500MB
【缺失：上传进度显示、失败重试机制】

2.2.2 文件共享功能
- 可设置共享链接有效期
- 可设置访问密码
【缺失：访问日志记录、下载次数限制】

## 3. 非功能需求
### 3.1 性能需求
3.1.1 响应时间
- 页面加载时间不超过3秒
【缺失：并发用户数指标、95%响应时间要求】

3.1.2 系统可用性
- 系统全年可用性不低于99.5%
【缺失：维护窗口定义、故障恢复时间】

### 3.2 安全需求
3.2.1 数据安全
- 敏感数据需加密存储
- 密码使用bcrypt加密
【缺失：加密算法标准、密钥管理机制】

## 4. 接口需求
### 4.1 邮件服务接口
4.1.1 发送邮件接口
- 支持发送HTML格式邮件
- 支持附件发送
【缺失：接口调用频率限制、错误码定义】

### 4.2 短信服务接口
4.2.1 发送短信接口
- 支持验证码短信
【完整缺失：接口地址、认证方式、参数格式】
"""
    
    # 示例2：电商系统需求
    ecommerce_content = """
# 电商平台需求规格说明书
版本：2.0
日期：2024-01-20

## 1. 功能需求
### 1.1 商品管理
1.1.1 商品上架
- 支持多规格商品
- 可设置库存预警值
要求：上架后需审核通过才可销售

1.1.2 商品搜索
- 支持关键词搜索
- 支持按价格、销量排序
要求：搜索结果响应时间不超过2秒

### 1.2 订单管理
1.2.1 下单流程
- 支持购物车下单
- 支持优惠券使用
要求：订单创建后30分钟内未支付自动取消

### 1.3 支付功能
1.3.1 支付方式
- 支持微信支付、支付宝
要求：支付成功率不低于99.9%

## 2. 非功能需求
### 2.1 性能需求
- 高峰时段支持10万并发用户
- API响应时间P95不超过200ms

### 2.2 安全需求
- 支付接口需支持防重放攻击
- 用户敏感信息需脱敏展示
"""
    
    # 示例3：政务系统需求
    government_content = """
# 政务服务系统需求规格说明书
版本：1.1
日期：2024-01-25

## 1. 功能需求
### 1.1 事项办理
1.1.1 在线申请
市民可在线上传材料申请事项办理。

1.1.2 进度查询
市民可实时查询办理进度。

### 1.2 材料管理
1.2.1 电子证照
支持调取电子证照库信息。

## 2. 非功能需求
### 2.1 可靠性
系统需保证7×24小时稳定运行。

### 2.2 安全性
需通过网络安全等级保护三级认证。
"""
    
    # 保存示例文档
    docs = [
        ("智能办公系统需求.docx", office_content),
        ("电商平台需求.docx", ecommerce_content),
        ("政务服务平台需求.docx", government_content)
    ]
    
    created_files = []
    for filename, content in docs:
        filepath = os.path.join("sample_docs", filename)
        doc = Document()
        doc.add_paragraph(content)
        doc.save(filepath)
        created_files.append(filepath)
        print(f"✓ 创建: {filename}")
    
    return created_files

def demo_document_preprocessing():
    """演示文档预处理"""
    print_header("文档预处理演示")
    
    # 初始化预处理器
    preprocessor = DocumentPreprocessor(max_segment_length=10000)
    
    # 处理示例文档
    sample_file = "sample_docs/智能办公系统需求.docx"
    if os.path.exists(sample_file):
        print(f"处理文档: {sample_file}")
        segments = preprocessor.process_document(sample_file)
        
        print(f"\n文档分割为 {len(segments)} 个片段:")
        for i, segment in enumerate(segments[:3], 1):  # 只显示前3个片段
            print(f"\n片段{i} (ID: {segment.id}):")
            preview = segment.text[:200].replace('\n', ' ')
            print(f"  预览: {preview}...")
        
        if len(segments) > 3:
            print(f"\n... 还有 {len(segments)-3} 个片段未显示")
    
    return preprocessor

def demo_single_document_validation(config):
    """演示单文档验证"""
    print_header("单文档验证演示")
    
    # 初始化验证器
    validator = RequirementValidator(config)
    
    # 验证智能办公系统需求
    sample_file = "sample_docs/智能办公系统需求.docx"
    if not os.path.exists(sample_file):
        print(f"文件不存在: {sample_file}")
        return None
    
    print(f"开始验证: {sample_file}")
    start_time = time.time()
    
    try:
        result = validator.validate_document(sample_file)
        elapsed_time = time.time() - start_time
        
        # 显示验证结果
        print(f"\n✓ 验证完成! 耗时: {elapsed_time:.2f}秒")
        print(f"\n验证结果摘要:")
        print(f"  文档名称: {result.document_name}")
        print(f"  完整性得分: {result.completeness_score:.2f}%")
        print(f"  总需求数: {result.total_requirements}")
        print(f"  完整需求数: {result.complete_requirements}")
        print(f"  验证时间: {result.generated_at}")
        
        # 显示缺失要素统计
        if result.missing_elements_by_type:
            print(f"\n缺失要素统计:")
            for req_type, missing_dict in result.missing_elements_by_type.items():
                if missing_dict:
                    print(f"  {req_type}:")
                    for element_type, count in missing_dict.items():
                        print(f"    - {element_type}: {count}处")
        else:
            print(f"\n✓ 未发现缺失要素")
        
        # 显示部分需求详情
        if result.requirements_details:
            print(f"\n需求详情 (显示前3个):")
            for i, req in enumerate(result.requirements_details[:3], 1):
                print(f"\n  {i}. 需求ID: {req.id}")
                print(f"     类型: {req.req_type.value}")
                print(f"     描述: {req.text[:80]}...")
                print(f"     完整性得分: {req.completeness_score:.1f}")
                if req.missing_elements:
                    print(f"     缺失要素: {', '.join(req.missing_elements)}")
                if req.improvement_suggestions:
                    print(f"     建议: {req.improvement_suggestions[0][:60]}...")
        
        return result
        
    except Exception as e:
        print(f"✗ 验证失败: {e}")
        return None

def demo_batch_validation(config):
    """演示批量验证"""
    print_header("批量验证演示")
    
    # 创建更多测试文档
    test_dir = "test_batch_docs"
    os.makedirs(test_dir, exist_ok=True)
    
    # 复制示例文档到测试目录
    sample_files = list(Path("sample_docs").glob("*.docx"))
    for sample_file in sample_files:
        shutil.copy(sample_file, os.path.join(test_dir, sample_file.name))
    
    print(f"创建测试目录: {test_dir}")
    print(f"包含 {len(sample_files)} 个测试文档")
    
    # 初始化验证器
    validator = RequirementValidator(config)
    
    # 批量验证
    print(f"\n开始批量验证...")
    start_time = time.time()
    
    try:
        results = validator.validate_batch(test_dir)
        elapsed_time = time.time() - start_time
        
        print(f"\n✓ 批量验证完成! 总耗时: {elapsed_time:.2f}秒")
        print(f"共处理 {len(results)} 个文档")
        
        if results:
            # 计算统计数据
            avg_score = sum(r.completeness_score for r in results) / len(results)
            min_score = min(r.completeness_score for r in results)
            max_score = max(r.completeness_score for r in results)
            total_reqs = sum(r.total_requirements for r in results)
            
            print(f"\n批量验证统计:")
            print(f"  平均完整性得分: {avg_score:.2f}%")
            print(f"  最高得分: {max_score:.2f}%")
            print(f"  最低得分: {min_score:.2f}%")
            print(f"  总需求数: {total_reqs}")
            
            # 显示各文档得分
            print(f"\n各文档得分:")
            for result in sorted(results, key=lambda x: x.completeness_score, reverse=True):
                completeness_status = "优秀" if result.completeness_score >= 80 else "良好" if result.completeness_score >= 60 else "需改进"
                print(f"  {result.document_name:30} {result.completeness_score:6.2f}% ({completeness_status})")
        
        return results
        
    except Exception as e:
        print(f"✗ 批量验证失败: {e}")
        return []

def demo_custom_configuration():
    """演示自定义配置"""
    print_header("自定义配置演示")
    
    # 创建自定义配置
    custom_config = Config()
    
    # 1. 修改API参数
    custom_config.TEMPERATURE = 0.1  # 更确定性的输出
    custom_config.MAX_TOKENS = 4096  # 更大的响应长度
    
    # 2. 自定义验证标准
    custom_config.COMPLETENESS_CRITERIA = {
        "功能需求": ["业务价值", "用户故事", "验收条件", "异常场景"],
        "非功能需求": ["性能指标", "安全要求", "兼容性", "可维护性"],
        "接口需求": ["协议规范", "数据格式", "错误处理", "版本管理"]
    }
    
    # 3. 修改路径配置
    custom_config.INPUT_DIR = "custom_input"
    custom_config.OUTPUT_DIR = "custom_reports"
    
    print("自定义配置已创建:")
    print(f"  - Temperature: {custom_config.TEMPERATURE}")
    print(f"  - Max Tokens: {custom_config.MAX_TOKENS}")
    print(f"  - 输入目录: {custom_config.INPUT_DIR}")
    print(f"  - 输出目录: {custom_config.OUTPUT_DIR}")
    print(f"\n自定义验证标准:")
    for req_type, elements in custom_config.COMPLETENESS_CRITERIA.items():
        print(f"  {req_type}: {', '.join(elements)}")
    
    return custom_config

def demo_error_handling(config):
    """演示错误处理"""
    print_header("错误处理演示")
    
    validator = RequirementValidator(config)
    
    # 测试1: 不存在的文件
    print("测试1: 验证不存在的文件")
    try:
        result = validator.validate_document("non_existent_file.docx")
        print("✗ 预期抛出异常但未抛出")
    except Exception as e:
        print(f"✓ 预期错误: {type(e).__name__}: {str(e)[:80]}...")
    
    # 测试2: 空文件
    print("\n测试2: 验证空文件")
    empty_file = "empty_test.docx"
    doc = Document()
    doc.save(empty_file)
    
    try:
        result = validator.validate_document(empty_file)
        if result.total_requirements == 0:
            print(f"✓ 空文件处理正常，未发现需求")
        else:
            print(f"✗ 空文件发现了 {result.total_requirements} 个需求")
    except Exception as e:
        print(f"处理空文件时出错: {e}")
    
    # 清理测试文件
    if os.path.exists(empty_file):
        os.remove(empty_file)
    
    # 测试3: 网络异常模拟（通过无效API密钥）
    print("\n测试3: 模拟API调用失败")
    invalid_config = Config()
    invalid_config.API_KEY = "invalid_key_123456"
    invalid_config.MAX_RETRIES = 1  # 减少重试次数
    
    try:
        invalid_validator = RequirementValidator(invalid_config)
        result = invalid_validator.validate_document("sample_docs/智能办公系统需求.docx")
        print("✗ 预期API调用失败但未失败")
    except Exception as e:
        print(f"✓ 预期API错误: {type(e).__name__}")
    
    print("\n✓ 错误处理演示完成")

def demo_report_generation(config, validation_result):
    """演示报告生成"""
    print_header("报告生成演示")
    
    if not validation_result:
        print("没有验证结果，跳过报告生成演示")
        return
    
    # 初始化报告生成器
    report_generator = ReportGenerator()
    
    # 生成Word报告
    word_report = "demo_validation_report.docx"
    print(f"生成Word报告: {word_report}")
    report_generator.generate_word_report(validation_result, word_report)
    
    # 生成Excel报告
    excel_report = "demo_validation_details.xlsx"
    print(f"生成Excel报告: {excel_report}")
    report_generator.generate_excel_report(validation_result, excel_report)
    
    # 生成JSON数据
    json_report = "demo_validation_data.json"
    print(f"生成JSON数据: {json_report}")
    
    # 将结果转换为字典
    result_dict = {
        "document_info": {
            "name": validation_result.document_name,
            "id": validation_result.document_id
        },
        "summary": {
            "total_requirements": validation_result.total_requirements,
            "complete_requirements": validation_result.complete_requirements,
            "completeness_score": validation_result.completeness_score,
            "validation_time": validation_result.validation_time
        },
        "requirements_summary": [
            {
                "id": req.id,
                "type": req.req_type.value,
                "score": req.completeness_score,
                "missing_count": len(req.missing_elements)
            }
            for req in validation_result.requirements_details[:10]  # 只保存前10个
        ]
    }
    
    with open(json_report, 'w', encoding='utf-8') as f:
        json.dump(result_dict, f, ensure_ascii=False, indent=2)
    
    print(f"\n✓ 报告生成完成!")
    print(f"  生成的报告:")
    print(f"  - {word_report}")
    print(f"  - {excel_report}")
    print(f"  - {json_report}")
    
    # 显示报告大小
    for report_file in [word_report, excel_report, json_report]:
        if os.path.exists(report_file):
            size = os.path.getsize(report_file) / 1024  # KB
            print(f"  {report_file:30} {size:.1f} KB")

def demo_advanced_features(config):
    """演示高级功能"""
    print_header("高级功能演示")
    
    # 1. 缓存功能演示
    print("1. 缓存功能演示")
    validator = RequirementValidator(config)
    
    sample_file = "sample_docs/智能办公系统需求.docx"
    
    if os.path.exists(sample_file):
        print(f"第一次验证 {sample_file} (会调用API)...")
        start_time = time.time()
        result1 = validator.validate_document(sample_file)
        time1 = time.time() - start_time
        
        print(f"第二次验证 {sample_file} (可能使用缓存)...")
        start_time = time.time()
        result2 = validator.validate_document(sample_file)
        time2 = time.time() - start_time
        
        print(f"\n验证时间对比:")
        print(f"  第一次: {time1:.2f}秒")
        print(f"  第二次: {time2:.2f}秒")
        
        if time2 < time1 * 0.8:  # 第二次快至少20%
            print(f"✓ 缓存生效，速度提升 {((time1-time2)/time1*100):.1f}%")
        else:
            print(f"注意: 缓存效果不明显")
    
    # 2. 并发处理演示
    print("\n2. 并发处理演示")
    validator.config.BATCH_SIZE = 3  # 增加并发数
    print(f"设置并发数: {validator.config.BATCH_SIZE}")
    
    # 3. 自定义验证标准应用
    print("\n3. 自定义验证标准应用")
    custom_criteria = {
        "功能需求": ["用户价值", "业务流程", "数据规则", "界面要求"],
        "非功能需求": ["SLA指标", "安全合规", "监控要求"],
        "接口需求": ["API文档", "版本策略", "限流策略"]
    }
    
    print("应用自定义标准:")
    for req_type, elements in custom_criteria.items():
        print(f"  {req_type}:")
        for element in elements:
            print(f"    - {element}")
    
    print("\n✓ 高级功能演示完成")

def main():
    """主函数"""
    print_header("需求完整性自动化验证系统演示")
    print("作者: 姚梦园")
    print("版本: 1.0")
    print("\n本演示展示如何使用模块化的需求完整性验证系统")
    
    # 检查API密钥
    api_key = os.getenv("DEEPSEEK_API_KEY")
    if not api_key or api_key == "your-api-key-here":
        print("\n⚠️  警告: 未检测到有效的DeepSeek API密钥")
        print("请设置环境变量: DEEPSEEK_API_KEY")
        print("或修改config.py中的API_KEY")
        
        use_demo = input("\n是否继续使用演示模式? (y/n): ").lower().strip()
        if use_demo != 'y':
            print("演示结束")
            return
    
    # 创建基本配置
    config = Config()
    if api_key and api_key != "your-api-key-here":
        config.API_KEY = api_key
    
    # 创建输出目录
    os.makedirs(config.OUTPUT_DIR, exist_ok=True)
    
    # 演示步骤
    steps = [
        ("创建示例文档", lambda: create_sample_documents()),
        ("文档预处理演示", lambda: demo_document_preprocessing()),
        ("单文档验证演示", lambda: demo_single_document_validation(config)),
        ("批量验证演示", lambda: demo_batch_validation(config)),
        ("自定义配置演示", demo_custom_configuration),
        ("错误处理演示", lambda: demo_error_handling(config)),
        ("高级功能演示", lambda: demo_advanced_features(config)),
    ]
    
    # 执行演示步骤
    validation_result = None
    
    for step_name, step_func in steps:
        try:
            print(f"\n{' 步骤开始 ':=^60}")
            print(f"正在执行: {step_name}")
            print('='*60)
            
            result = step_func()
            if step_name == "单文档验证演示":
                validation_result = result
            
            print(f"\n✓ {step_name} 完成")
            
        except Exception as e:
            print(f"\n✗ {step_name} 失败: {e}")
            import traceback
            traceback.print_exc()
        
        # 询问是否继续
        if step_name != steps[-1][0]:  # 不是最后一步
            cont = input(f"\n按Enter继续下一步，或输入q退出演示: ").strip()
            if cont.lower() == 'q':
                print("用户退出演示")
                break
    
    # 最后演示报告生成（如果有验证结果）
    if validation_result:
        print(f"\n{' 报告生成 ':=^60}")
        demo_report_generation(config, validation_result)
    
    # 清理演示文件
    cleanup = input("\n是否清理演示生成的文件? (y/n): ").lower().strip()
    if cleanup == 'y':
        print("清理演示文件...")
        for dir_name in ["sample_docs", "test_batch_docs", "custom_input", "custom_reports", "cache"]:
            if os.path.exists(dir_name):
                shutil.rmtree(dir_name)
                print(f"  删除目录: {dir_name}")
        
        for file_name in ["demo_validation_report.docx", "demo_validation_details.xlsx", 
                         "demo_validation_data.json", "empty_test.docx"]:
            if os.path.exists(file_name):
                os.remove(file_name)
                print(f"  删除文件: {file_name}")
        
        print("✓ 清理完成")
    
    print_header("演示完成")
    print("感谢使用需求完整性自动化验证系统!")
    print(f"生成的报告可在 '{config.OUTPUT_DIR}' 目录中找到")

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n演示被用户中断")
    except Exception as e:
        print(f"\n演示发生错误: {e}")
        import traceback
        traceback.print_exc()