# ==================== preprocessor.py ====================
import re
from pathlib import Path
from typing import List
import docx
from docx import Document
import PyPDF2

from models import DocumentSegment

class TextCleaner:
    """文本清理工具类"""
    
    def __init__(self):
        self.header_footer_keywords = [
            '机密', '保密', '第\d+页', '共\d+页',
            'copyright', '©', 'confidential',
            'header', 'footer', '页眉', '页脚'
        ]
        
    def clean(self, text: str) -> str:
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if not line or len(line) < 3:
                continue
            if self._is_header_footer(line):
                continue
            line = re.sub(r'\s+', ' ', line)
            cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def _is_header_footer(self, line: str) -> bool:
        line_lower = line.lower()
        if len(line) < 50:
            for keyword in self.header_footer_keywords:
                if re.search(keyword, line_lower, re.IGNORECASE):
                    return True
        if re.match(r'^[\-\s]*\d+[\-\s]*$', line):
            return True
        return False

class DocumentPreprocessor:
    """文档预处理类"""
    
    def __init__(self, max_segment_length: int = 30000):
        self.max_segment_length = max_segment_length
        self.text_cleaner = TextCleaner()
        
    def process_document(self, file_path: str) -> List[DocumentSegment]:
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            text = self._extract_from_pdf(file_path)
        elif file_ext in ['.doc', '.docx']:
            text = self._extract_from_word(file_path)
        elif file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        else:
            raise ValueError(f"不支持的文件格式: {file_ext}")
        
        cleaned_text = self.text_cleaner.clean(text)
        segments = self._split_text(cleaned_text, Path(file_path).name)
        return segments
    
    def _extract_from_pdf(self, file_path: str) -> str:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                page_text = re.sub(r'\s+', ' ', page_text)
                page_text = page_text.replace('-\n', '')
                text += f"\n--- 第 {page_num + 1} 页 ---\n{page_text}"
        return text
    
    def _extract_from_word(self, file_path: str) -> str:
        text = ""
        doc = Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"
        return text
    
    def _split_text(self, text: str, filename: str) -> List[DocumentSegment]:
        # 尝试按章节分割
        chapter_patterns = [
            r'第[一二三四五六七八九十\d]+章\s+[^\n]+',
            r'\d+\.\d+\s+[^\n]+',
            r'[A-Z]\.\d+\s+[^\n]+'
        ]
        
        chapter_positions = []
        for pattern in chapter_patterns:
            for match in re.finditer(pattern, text):
                chapter_positions.append(match.start())
        
        if not chapter_positions:
            return self._split_by_length(text, filename)
        
        chapter_positions.sort()
        chapter_positions.append(len(text))
        
        segments = []
        for i in range(len(chapter_positions) - 1):
            start = chapter_positions[i]
            end = chapter_positions[i + 1]
            segment_text = text[start:end].strip()
            if segment_text:
                segment_id = f"{filename}_ch{len(segments)+1}"
                segments.append(DocumentSegment(
                    id=segment_id,
                    text=segment_text,
                    original_file=filename
                ))
        
        final_segments = []
        for segment in segments:
            if len(segment.text) > self.max_segment_length:
                sub_segments = self._split_by_length(segment.text, segment.id)
                final_segments.extend(sub_segments)
            else:
                final_segments.append(segment)
        
        return final_segments
    
    def _split_by_length(self, text: str, base_id: str) -> List[DocumentSegment]:
        segments = []
        words = text.split()
        current_segment = []
        current_length = 0
        segment_num = 1
        
        for word in words:
            word_length = len(word) + 1
            if current_length + word_length > self.max_segment_length:
                segment_text = ' '.join(current_segment)
                segments.append(DocumentSegment(
                    id=f"{base_id}_part{segment_num}",
                    text=segment_text,
                    original_file=base_id
                ))
                current_segment = [word]
                current_length = word_length
                segment_num += 1
            else:
                current_segment.append(word)
                current_length += word_length
        
        if current_segment:
            segment_text = ' '.join(current_segment)
            segments.append(DocumentSegment(
                id=f"{base_id}_part{segment_num}",
                text=segment_text,
                original_file=base_id
            ))
        
        return segments